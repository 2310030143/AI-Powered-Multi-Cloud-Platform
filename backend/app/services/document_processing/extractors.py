"""Text extraction for supported file types.

Every extractor returns an :class:`ExtractionResult` with per-page text so the
chunking stage can keep page metadata (where the source supports pages).
"""
import io
from dataclasses import dataclass, field

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.config.settings import get_settings

settings = get_settings()

IMAGE_TYPES = {"png", "jpg", "jpeg", "tiff"}


@dataclass
class PageText:
    page_number: int  # 1-based
    text: str


@dataclass
class ExtractionResult:
    text: str = ""
    pages: list[PageText] = field(default_factory=list)
    # True when one or more pages need OCR
    ocr_recommended: bool = False
    # 1-based page numbers that need OCR
    ocr_page_numbers: list[int] = field(default_factory=list)
    # Populated for CSV sources: rows as list of lists of strings
    csv_rows: list[list[str]] | None = None


class UnsupportedFileType(Exception):
    """Raised when no extractor exists for a file type."""


class ExtractionFailed(Exception):
    """Raised when the file exists but cannot be parsed."""


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def extract_pdf(content: bytes) -> ExtractionResult:
    """Extract text per page with pypdf; flags scanned-looking PDFs."""
    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            reader.decrypt("")  # try empty user password
    except (PdfReadError, ValueError) as exc:
        raise ExtractionFailed(f"Could not read PDF: {exc}") from exc

    pages: list[PageText] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = _sanitize_text(page.extract_text() or "").strip()
        except Exception:  # a single broken page should not kill extraction
            text = ""
        pages.append(PageText(page_number=index, text=text))

    full_text = "\n\n".join(p.text for p in pages if p.text)

    ocr_page_numbers = [
        page.page_number
        for page in pages
        if not page.text.strip()
    ]

    return ExtractionResult(
        text=full_text,
        pages=pages,
        ocr_recommended=bool(ocr_page_numbers),
        ocr_page_numbers=ocr_page_numbers,
    )


def extract_docx(content: bytes) -> ExtractionResult:
    """Extract paragraphs and tables from a .docx file."""
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionFailed(f"Could not read DOCX: {exc}") from exc

    parts: list[str] = [
    _sanitize_text(p.text).strip()
    for p in document.paragraphs
    if p.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [_sanitize_text(c.text).strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts)
    return ExtractionResult(text=text, pages=[PageText(1, text)] if text else [])


def extract_txt(content: bytes) -> ExtractionResult:
    text = _sanitize_text(_decode_text(content)).strip()
    return ExtractionResult(text=text, pages=[PageText(1, text)] if text else [])


def extract_csv(content: bytes) -> ExtractionResult:
    """Parse CSV into text lines (``column: value`` per row) plus raw rows."""
    raw = io.BytesIO(content)
    try:
        df = pd.read_csv(raw, dtype=str, keep_default_na=False, encoding="utf-8")
    except UnicodeDecodeError:
        raw.seek(0)
        df = pd.read_csv(raw, dtype=str, keep_default_na=False, encoding="latin-1")
    except pd.errors.EmptyDataError as exc:
        raise ExtractionFailed("CSV file has no data") from exc
    except pd.errors.ParserError as exc:
        raise ExtractionFailed(f"Could not parse CSV: {exc}") from exc

    columns = [str(c) for c in df.columns]
    rows = [
        [("" if v is None else _sanitize_text(str(v))) for v in row]
        for row in df.values.tolist()
    ]
    lines = []
    for row in rows:
        pairs = [f"{col}: {val}" for col, val in zip(columns, row) if str(val).strip()]
        if pairs:
            lines.append(", ".join(pairs))
    text = "\n\n".join(lines)
    return ExtractionResult(
        text=text,
        pages=[PageText(1, text)] if text else [],
        csv_rows=[columns] + rows if rows else None,
    )


def extract_image(content: bytes) -> ExtractionResult:
    """Images carry no extractable text — OCR is always recommended."""
    return ExtractionResult(text="", pages=[], ocr_recommended=True)


def extract(content: bytes, file_type: str) -> ExtractionResult:
    """Dispatch extraction by file type/extension."""
    file_type = (file_type or "").lower().lstrip(".")
    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "txt": extract_txt,
        "csv": extract_csv,
    }
    if file_type in IMAGE_TYPES:
        return extract_image(content)
    extractor = extractors.get(file_type)
    if extractor is None:
        raise UnsupportedFileType(
            f"No extractor for file type '{file_type}'. Supported: "
            f"{', '.join(sorted(extractors))}, {', '.join(sorted(IMAGE_TYPES))} (via OCR)"
        )
    return extractor(content)


def _sanitize_text(text: str) -> str:
    """Remove NUL characters that PostgreSQL cannot store in text fields."""
    return text.replace("\x00", "")

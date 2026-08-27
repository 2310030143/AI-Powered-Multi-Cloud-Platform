"""Shared helpers for Phase 3 tests."""


def make_pdf(texts: list[str]) -> bytes:
    """Build a minimal valid PDF with one page per text string.

    Hand-rolled so tests need no extra dependencies; pypdf parses the result.
    """
    n_pages = len(texts)
    page_ids = [3 + 2 * i for i in range(n_pages)]  # 3,5,7,...
    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")

    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>".encode())

    for i, text in enumerate(texts):
        page_id = page_ids[i]
        content_id = page_id + 1
        safe = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({safe}) Tj ET".encode()
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents {content_id} 0 R "
            f"/Resources << /Font << /F1 {2 + 2 * n_pages + 1} 0 R >> >> >>".encode()
        )
        objects.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )

    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"

    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF"
    ).encode()
    return bytes(out)


def make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build a DOCX in memory with python-docx."""
    import io
    from docx import Document

    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell in enumerate(row):
                table.rows[r].cells[c].text = cell
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
